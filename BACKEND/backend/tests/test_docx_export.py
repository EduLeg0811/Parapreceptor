from __future__ import annotations

import os
import unittest
from io import BytesIO
from unittest.mock import patch

from docx import Document
from fastapi.testclient import TestClient

from backend.functions.docx_export import build_lexical_overview_docx, build_semantic_overview_docx
from backend.main import app


class DocxExportTests(unittest.TestCase):
    def test_build_lexical_overview_docx_contains_title_metadata_and_result(self) -> None:
        content = build_lexical_overview_docx({
            "term": "cosmoetica",
            "totalFound": 1,
            "groups": [
                {
                    "bookCode": "LO",
                    "bookLabel": "Lexico de Ortopensatas",
                    "totalFound": 1,
                    "matches": [
                        {
                            "text": "Trecho com **cosmoetica**. (p. 41)",
                            "title": "Cosmoetica",
                            "pagina": "41",
                            "row": 1,
                        }
                    ],
                }
            ],
        })

        document = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn("Lexical Overview", text)
        self.assertIn("Termo: cosmoetica", text)
        self.assertIn("Lexico de Ortopensatas", text)
        self.assertIn("Trecho com cosmoetica.", text)
        self.assertIn("Lexico de Ortopensatas | Cosmoetica | p. 41", text)

        self.assertIn("1. Trecho com cosmoetica.", [paragraph.text for paragraph in document.paragraphs])
        self.assertNotIn("1. Trecho com cosmoetica. (p. 41)", [paragraph.text for paragraph in document.paragraphs])

        metadata_paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text == "Lexico de Ortopensatas | Cosmoetica | p. 41")
        self.assertEqual(metadata_paragraph.paragraph_format.space_after.pt, 6)

        group_index = next(index for index, paragraph in enumerate(document.paragraphs) if paragraph.text == "Lexico de Ortopensatas")
        self.assertEqual(document.paragraphs[group_index + 1].text, "")

    def test_build_semantic_overview_docx_contains_score_and_safety_note(self) -> None:
        content = build_semantic_overview_docx({
            "term": "cosmoetica",
            "minScore": 0.6,
            "totalIndexes": 1,
            "totalFound": 1,
            "lexicalFilteredCount": 0,
            "safetyNote": "Exportacao ALL seguro.",
            "groups": [
                {
                    "indexId": "lo",
                    "indexLabel": "LO Semantic",
                    "totalFound": 1,
                    "matches": [
                        {
                            "book": "LO",
                            "row": 2,
                            "text": "Resultado semanticamente afim.",
                            "metadata": {"title": "Cosmoetica", "pagina": "41"},
                            "score": 0.9876,
                        }
                    ],
                }
            ],
        })

        document = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertIn("Semantic Overview", text)
        self.assertIn("Score minimo: 0.60", text)
        self.assertIn("Exportacao ALL seguro.", text)
        self.assertIn("LO | Cosmoetica | p. 41 | Score: 0.9876", text)
        self.assertNotIn("Livro/Base:", text)
        self.assertNotIn("Titulo:", text)
        self.assertNotIn("Pagina:", text)
        self.assertNotIn("Registro:", text)

        metadata_paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text == "LO | Cosmoetica | p. 41 | Score: 0.9876")
        self.assertEqual(metadata_paragraph.paragraph_format.space_after.pt, 6)

        group_index = next(index for index, paragraph in enumerate(document.paragraphs) if paragraph.text == "LO Semantic")
        self.assertEqual(document.paragraphs[group_index + 1].text, "")


class DocxExportApiTests(unittest.TestCase):
    def setUp(self) -> None:
        self.client = TestClient(app)

    @patch("backend.functions.lexical_search_service.search_lexical_overview_with_total")
    def test_lexical_overview_export_endpoint_returns_docx_download(self, mock_search) -> None:
        mock_search.return_value = (
            1,
            [
                {
                    "bookCode": "LO",
                    "bookLabel": "Lexico de Ortopensatas",
                    "fileStem": "LO",
                    "totalFound": 1,
                    "shownCount": 1,
                    "matches": [{"text": "Trecho lexical.", "row": 1, "pagina": "12"}],
                }
            ],
        )

        response = self.client.post(
            "/api/apps/lexical/overview/export-docx",
            json={"term": "cosmoetica", "miniTextWindow": 3, "maxResultsDocx": 123},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(mock_search.call_args.kwargs["limit"], 200)
        self.assertEqual(response.headers["content-type"], "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertIn("attachment", response.headers["content-disposition"])
        self.assertGreater(len(response.content), 1000)

    @patch.dict(os.environ, {"OPENAI_API_KEY": "test-key"})
    @patch("backend.functions.semantic_search_service.search_semantic_overview_export_with_total")
    def test_semantic_overview_export_endpoint_returns_docx_download(self, mock_search) -> None:
        mock_search.return_value = (
            1,
            1,
            0,
            0.6,
            0.6,
            {"usedRagContext": False},
            [
                {
                    "indexId": "lo",
                    "indexLabel": "LO Semantic",
                    "totalFound": 1,
                    "shownCount": 1,
                    "matches": [{"book": "LO", "row": 1, "text": "Trecho semantic.", "metadata": {}, "score": 0.9}],
                }
            ],
            {"perIndexLimit": 200, "totalLimit": 123, "includedCount": 1, "truncated": False},
        )

        response = self.client.post(
            "/api/apps/semantic/overview/export-docx",
            json={"term": "cosmoetica", "minScore": 0.6, "miniTextWindow": 3, "maxResultsDocx": 123},
        )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(mock_search.call_args.kwargs["total_limit"], 123)
        self.assertEqual(response.headers["content-type"], "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertIn("attachment", response.headers["content-disposition"])
        self.assertGreater(len(response.content), 1000)

    @patch("backend.functions.lexical_search_service.search_lexical_overview_with_total")
    def test_lexical_overview_export_from_payload_does_not_rerun_search(self, mock_search) -> None:
        response = self.client.post(
            "/api/apps/lexical/overview/export-docx-from-payload",
            json={
                "term": "cosmoetica",
                "totalBooks": 1,
                "totalFound": 1,
                "maxResultsDocx": 123,
                "groups": [
                    {
                        "bookCode": "LO",
                        "bookLabel": "Lexico de Ortopensatas",
                        "fileStem": "LO",
                        "totalFound": 1,
                        "shownCount": 1,
                        "matches": [{"text": "Trecho lexical.", "row": 1, "pagina": "12"}],
                    }
                ],
            },
        )

        self.assertEqual(response.status_code, 200)
        mock_search.assert_not_called()
        self.assertEqual(response.headers["content-type"], "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertIn("attachment", response.headers["content-disposition"])
        self.assertGreater(len(response.content), 1000)

    @patch.dict(os.environ, {"OPENAI_API_KEY": "test-key"})
    @patch("backend.functions.semantic_search_service.search_semantic_overview_export_with_total")
    def test_semantic_overview_export_from_payload_does_not_rerun_search(self, mock_search) -> None:
        response = self.client.post(
            "/api/apps/semantic/overview/export-docx-from-payload",
            json={
                "term": "cosmoetica",
                "minScore": 0.6,
                "recommendedMinScoreMin": 0.6,
                "recommendedMinScoreMax": 0.6,
                "usesCalibratedMinScores": False,
                "ignoreBaseCalibration": True,
                "excludeLexicalDuplicates": False,
                "totalIndexes": 1,
                "totalFound": 1,
                "lexicalFilteredCount": 0,
                "maxResultsDocx": 123,
                "groups": [
                    {
                        "indexId": "lo",
                        "indexLabel": "LO Semantic",
                        "totalFound": 1,
                        "shownCount": 1,
                        "matches": [{"book": "LO", "row": 1, "text": "Trecho semantic.", "metadata": {}, "score": 0.9}],
                    }
                ],
            },
        )

        self.assertEqual(response.status_code, 200)
        mock_search.assert_not_called()
        self.assertEqual(response.headers["content-type"], "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertIn("attachment", response.headers["content-disposition"])
        self.assertGreater(len(response.content), 1000)


if __name__ == "__main__":
    unittest.main()
