from __future__ import annotations

import re
from datetime import datetime, timezone
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

WORD_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MAX_RESULTS_DOCX = 200


def _clean_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").replace("\u00a0", " ")).strip()


def _safe_filename_part(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_-]+", "-", _clean_text(value))[:80].strip("-")
    return cleaned or "overview"


def build_docx_filename(prefix: str, term: str) -> str:
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
    return f"{prefix}-{_safe_filename_part(term)}-{stamp}.docx"


def _set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 1


def _add_paragraph_border(paragraph: Any, *, bottom_only: bool = False, style: str = "single", color: str = "8FAADC") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)

    sides = ("bottom",) if bottom_only else ("top", "bottom")
    for side in sides:
        element = borders.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            borders.append(element)
        element.set(qn("w:val"), style)
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "3")
        element.set(qn("w:color"), color)


def _add_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    _add_paragraph_border(paragraph, style="double", color="4F81BD")

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _add_section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x80, 0x40)
    _add_paragraph_border(paragraph, bottom_only=True, color="70AD47")


def _add_blank_line(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)


def _add_metadata_line(document: Document, label: str, value: Any) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1)
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(9)
    value_run = paragraph.add_run(_clean_text(value))
    value_run.font.size = Pt(9)


def _add_stats(document: Document, stats: list[tuple[str, Any]]) -> None:
    _add_section_heading(document, "Estatisticas")
    for label, value in stats:
        _add_metadata_line(document, label, value)


def _add_markdown_runs(paragraph: Any, text: str) -> None:
    source = _clean_text(text)
    if not source:
        paragraph.add_run("")
        return

    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(source):
        if match.start() > pos:
            paragraph.add_run(source[pos:match.start()])
        token = match.group(0)
        run = paragraph.add_run(token.strip("*"))
        if token.startswith("**"):
            run.bold = True
        else:
            run.italic = True
        pos = match.end()
    if pos < len(source):
        paragraph.add_run(source[pos:])


def _strip_trailing_page_marker(text: Any, page: str) -> str:
    cleaned_text = _clean_text(text)
    cleaned_page = _clean_text(page)
    if not cleaned_text or not cleaned_page:
        return cleaned_text
    return re.sub(rf"\s*\(p\.\s*{re.escape(cleaned_page)}\)\s*$", "", cleaned_text).strip()


def _add_result(document: Document, number: int, text: str, metadata: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(4)
    prefix = paragraph.add_run(f"{number}. ")
    prefix.bold = True
    prefix.font.size = Pt(9)
    prefix.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    _add_markdown_runs(paragraph, text)

    if metadata:
        meta = document.add_paragraph()
        meta.paragraph_format.space_after = Pt(6)
        run = meta.add_run(metadata)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x20, 0x20)
    else:
        paragraph.paragraph_format.space_after = Pt(6)


def _docx_bytes(document: Document) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_lexical_overview_docx(payload: dict[str, Any]) -> bytes:
    document = Document()
    _set_document_defaults(document)

    term = _clean_text(payload.get("term"))
    groups = [group for group in payload.get("groups") or [] if isinstance(group, dict)]
    total_found = int(payload.get("totalFound") or 0)
    generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")

    _add_title(document, "Lexical Overview", f"Termo: {term}")
    _add_section_heading(document, "Dados da busca")
    _add_metadata_line(document, "Termo", term)
    _add_metadata_line(document, "Total encontrado", total_found)
    _add_metadata_line(document, "Livros com resultado", len(groups))
    _add_metadata_line(document, "Gerado em", generated_at)
    if payload.get("safetyNote"):
        _add_metadata_line(document, "Nota", payload.get("safetyNote"))

    _add_stats(
        document,
        [
            (_clean_text(group.get("bookLabel") or group.get("bookCode")), int(group.get("totalFound") or 0))
            for group in groups
        ],
    )

    _add_section_heading(document, "Resultados")
    counter = 1
    for group in groups:
        title = _clean_text(group.get("bookLabel") or group.get("bookCode") or "Livro")
        book_label = title
        _add_blank_line(document)
        _add_section_heading(document, title)
        _add_blank_line(document)
        for match in group.get("matches") or []:
            if not isinstance(match, dict):
                continue
            page = _clean_text(match.get("pagina"))
            text = _strip_trailing_page_marker(match.get("text"), page)
            item_title = _clean_text(match.get("title"))
            parts = [part for part in [book_label, item_title, f"p. {page}" if page else ""] if part]
            _add_result(document, counter, text, " | ".join(parts))

            #adicionar espaçamento entre resultados
            #_add_blank_line(document)
            counter += 1

    return _docx_bytes(document)


def build_semantic_overview_docx(payload: dict[str, Any]) -> bytes:
    document = Document()
    _set_document_defaults(document)

    term = _clean_text(payload.get("term"))
    groups = [group for group in payload.get("groups") or [] if isinstance(group, dict)]
    total_found = int(payload.get("totalFound") or 0)
    lexical_filtered = int(payload.get("lexicalFilteredCount") or 0)
    generated_at = datetime.now().strftime("%d/%m/%Y %H:%M")

    _add_title(document, "Semantic Overview", f"Query: {term}")
    _add_section_heading(document, "Dados da busca")
    _add_metadata_line(document, "Query", term)
    _add_metadata_line(document, "Total semantic", total_found)
    _add_metadata_line(document, "Bases analisadas", payload.get("totalIndexes") or len(groups))
    _add_metadata_line(document, "Score minimo", f"{float(payload.get('minScore') or 0):.2f}")
    _add_metadata_line(document, "Duplicados lexicos filtrados", lexical_filtered)
    _add_metadata_line(document, "Gerado em", generated_at)
    if payload.get("safetyNote"):
        _add_metadata_line(document, "Nota", payload.get("safetyNote"))

    _add_stats(
        document,
        [
            (_clean_text(group.get("indexLabel") or group.get("indexId")), int(group.get("totalFound") or 0))
            for group in groups
        ],
    )

    _add_section_heading(document, "Resultados")
    counter = 1
    for group in groups:
        title = _clean_text(group.get("indexLabel") or group.get("indexId") or "Base")
        _add_blank_line(document)
        _add_section_heading(document, title)
        _add_blank_line(document)
        for match in group.get("matches") or []:
            if not isinstance(match, dict):
                continue
            score = float(match.get("score") or 0.0)
            metadata = match.get("metadata") if isinstance(match.get("metadata"), dict) else {}
            item_title = _clean_text(metadata.get("title") or metadata.get("titulo"))
            page = _clean_text(metadata.get("pagina") or metadata.get("page"))
            parts = [
                _clean_text(match.get("book") or group.get("indexLabel")),
                item_title,
                f"p. {page}" if page else "",
                f"Score: {score:.4f}",
            ]
            _add_result(document, counter, _clean_text(match.get("text")), " | ".join(part for part in parts if part))
            
            
            counter += 1

    return _docx_bytes(document)
